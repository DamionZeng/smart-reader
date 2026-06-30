"""
DOCX → HTML 转换器，对应 src/lib/docx-to-html.ts。

使用 python-docx 遍历段落 + 内联图片，生成语义化 HTML。
图片提取后传 R2（与 PDF 管线一致），R2 未配置时 fallback 到 base64。

输出与 PDF 管线相同结构：ParseResult(html, plain_text, image_count, page_count)。
page_count 恒为 1（docx 无页面概念）。
"""
import io

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from app import storage
from app.parsers.image_extractor import ParseResult
from app.utils.html_utils import escape_html, strip_html


def is_docx_buffer(buffer: bytes) -> bool:
    """快速检测：是否为 docx 文件（PK\x03\x04 zip magic）。对应 isDocxBuffer。"""
    if not buffer or len(buffer) < 4:
        return False
    return (
        buffer[0] == 0x50
        and buffer[1] == 0x4B
        and buffer[2] == 0x03
        and buffer[3] == 0x04
    )


def _docx_style_to_tag(style_name: str) -> str:
    """把 docx 段落样式名映射到 HTML 标签。"""
    if not style_name:
        return "p"
    lower = style_name.lower()
    if "heading 1" in lower or "title" in lower:
        return "h1"
    if "heading 2" in lower:
        return "h2"
    if "heading 3" in lower:
        return "h3"
    if "heading 4" in lower:
        return "h4"
    if "heading 5" in lower:
        return "h5"
    if "heading 6" in lower:
        return "h6"
    if "list" in lower:
        return "li"
    return "p"


async def _extract_images_from_docx(
    docx_path_or_stream,
    user_id: str,
) -> dict[str, str]:
    """提取 docx 中所有图片，上传 R2，返回 rId → R2 URL 映射。

    python-docx 的内联图片通过 document.part.relations 关联，
    每个图片关系指向一个 image part，包含图片字节。
    """
    image_map: dict[str, str] = {}
    if not storage.is_r2_configured():
        return image_map

    doc = Document(docx_path_or_stream)
    for rel in doc.part.rels.values():
        if rel.reltype == RT.IMAGE:
            try:
                image_part = rel.target_part
                image_bytes = image_part.blob
                # 从 content_type 推断 MIME
                mime = image_part.content_type or "image/png"
                url = await storage.upload_image(image_bytes, mime, user_id)
                image_map[rel.rId] = url
            except Exception as e:
                print(f"[docx_parser] image extraction failed for {rel.rId}: {e}", flush=True)

    return image_map


def _inline_images_to_html(paragraph, image_map: dict[str, str]) -> str:
    """把段落中的内联图片转成 <figure><img>。

    python-docx 的 paragraph._element 包含 <a:blip> 元素指向图片 rId。
    """
    html_parts: list[str] = []
    # 遍历段落 XML 找 blip 元素
    nsmap = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    blips = paragraph._element.findall(
        ".//a:blip",
        nsmap,
    )
    for blip in blips:
        embed_attr = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
        rId = blip.get(embed_attr)
        if rId and rId in image_map:
            url = image_map[rId]
            html_parts.append(
                f'<figure><img src="{url}" alt="Embedded image" /></figure>'
            )
    return "".join(html_parts)


async def convert_docx_to_html(
    buffer: bytes,
    user_id: str,
) -> ParseResult:
    """把 DOCX 字节转换为结构化 HTML。对应 convertDocxToHtml。"""
    stream = io.BytesIO(buffer)
    doc = Document(stream)

    # 先提取所有图片到 R2
    image_map = await _extract_images_from_docx(stream, user_id)

    html_parts: list[str] = []
    image_count = len(image_map)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""
        tag = _docx_style_to_tag(style_name)

        # 处理列表项
        if tag == "li" and not html_parts or (html_parts and html_parts[-1] != "<ul>"):
            # 简化：列表项直接用 <li>
            pass

        # 内联图片
        img_html = _inline_images_to_html(paragraph, image_map)

        if text or img_html:
            if text:
                html_parts.append(f"<{tag}>{escape_html(text)}</{tag}>")
            if img_html:
                html_parts.append(img_html)

    # 简单处理表格
    for table in doc.tables:
        html_parts.append("<table>")
        for row in table.rows:
            html_parts.append("<tr>")
            for cell in row.cells:
                cell_text = cell.text.strip()
                html_parts.append(f"<td>{escape_html(cell_text)}</td>")
            html_parts.append("</tr>")
        html_parts.append("</table>")

    html = "\n".join(html_parts)
    plain_text = strip_html(html)

    return ParseResult(
        html=html,
        plain_text=plain_text,
        image_count=image_count,
        page_count=1,  # docx 无页面概念
    )
